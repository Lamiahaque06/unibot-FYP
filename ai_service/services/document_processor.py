"""
Document processing service.

Handles:
  - PDF text extraction
  - Smart recursive text chunking (respects sentence/paragraph boundaries)
  - Chunk metadata enrichment
"""

import re
import uuid
from pathlib import Path
from typing import List, Optional, Tuple

from models.schemas import DocumentChunk
from utils.logger import get_logger

logger = get_logger("document_processor")


# ─── Text Extraction ──────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> Tuple[str, List[str]]:
    """
    Extract text from a PDF file.
    Returns (full_text, per_page_texts).
    """
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text)
        full_text = "\n\n".join(pages)
        logger.info(f"Extracted {len(pages)} pages, {len(full_text)} chars from {file_path}")
        return full_text, pages
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise


def extract_text_from_txt(file_path: str) -> Tuple[str, List[str]]:
    """Extract text from a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return text, [text]


def extract_text_from_docx(file_path: str) -> Tuple[str, List[str]]:
    """
    Extract text from a .docx file using python-docx.

    Captures:
      - Body paragraphs (preserving heading structure with markers)
      - All table cells (row-by-row, pipe-separated)
      - Headers and footers from each section

    Returns (full_text, logical_sections) where sections approximate
    the document's section breaks (used as stand-ins for page numbers).
    """
    try:
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn
    except ImportError:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    doc = DocxDocument(file_path)
    sections_text: List[str] = []
    current_section: List[str] = []

    def flush_section() -> None:
        text = "\n".join(l for l in current_section if l.strip())
        if text:
            sections_text.append(text)
        current_section.clear()

    # ── Body paragraphs ──────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower()

        # Use heading styles as section markers
        if style_name.startswith("heading"):
            flush_section()
            level = style_name.replace("heading", "").strip()
            prefix = "#" * (int(level) if level.isdigit() else 1)
            current_section.append(f"{prefix} {text}")
        else:
            current_section.append(text)

    flush_section()

    # ── Tables ───────────────────────────────────────────────
    table_blocks: List[str] = []
    for table in doc.tables:
        rows: List[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Deduplicate merged cells
            seen = []
            for c in cells:
                if not seen or c != seen[-1]:
                    seen.append(c)
            rows.append(" | ".join(c for c in seen if c))
        table_text = "\n".join(r for r in rows if r)
        if table_text:
            table_blocks.append(table_text)

    if table_blocks:
        sections_text.append("Tables:\n" + "\n\n".join(table_blocks))

    # ── Headers / Footers ────────────────────────────────────
    hf_lines: List[str] = []
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None:
                for para in hf.paragraphs:
                    t = para.text.strip()
                    if t:
                        hf_lines.append(t)

    if hf_lines:
        # Deduplicate repeated header/footer lines
        unique_hf = list(dict.fromkeys(hf_lines))
        sections_text.insert(0, "Document header/footer: " + " | ".join(unique_hf))

    full_text = "\n\n".join(sections_text)
    logger.info(
        f"Extracted {len(sections_text)} logical sections, "
        f"{len(full_text)} chars from DOCX {file_path}"
    )
    return full_text, sections_text


def extract_text_from_txt(file_path: str) -> Tuple[str, List[str]]:
    """Extract text from a plain text file."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return text, [text]


def extract_text(file_path: str) -> Tuple[str, List[str]]:
    """Auto-detect file type and extract text."""
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx",):
        return extract_text_from_docx(file_path)
    elif ext in (".txt", ".md"):
        return extract_text_from_txt(file_path)
    elif ext == ".doc":
        raise ValueError(
            "Legacy .doc format is not supported. "
            "Please convert to .docx (File → Save As in Word) and re-upload."
        )
    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ─── Text Cleaning ────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """Normalize whitespace while preserving paragraph structure."""
    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Collapse multiple spaces
    text = re.sub(r" {2,}", " ", text)
    # Remove form-feed characters
    text = text.replace("\f", "\n\n")
    return text.strip()


# ─── Recursive Chunking ───────────────────────────────────────────────────────

SEPARATORS = ["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " "]


def _split_by_separator(text: str, separator: str, chunk_size: int) -> List[str]:
    """Split text by a given separator if chunks would be too large."""
    parts = text.split(separator)
    chunks = []
    current = ""
    for part in parts:
        candidate = (current + separator + part) if current else part
        if len(candidate) <= chunk_size:
            current = candidate
        else:
            if current:
                chunks.append(current.strip())
            current = part
    if current:
        chunks.append(current.strip())
    return [c for c in chunks if c]


def recursive_chunk(
    text: str,
    chunk_size: int = 800,
    chunk_overlap: int = 120,
    separators: Optional[List[str]] = None,
) -> List[str]:
    """
    Recursively split text using a hierarchy of separators.
    Similar to LangChain's RecursiveCharacterTextSplitter but dependency-free.
    """
    if separators is None:
        separators = SEPARATORS

    if len(text) <= chunk_size:
        return [text] if text.strip() else []

    # Try each separator level
    for sep in separators:
        if sep in text:
            parts = _split_by_separator(text, sep, chunk_size)
            if len(parts) > 1:
                break
    else:
        # Hard split as last resort
        parts = [text[i : i + chunk_size] for i in range(0, len(text), chunk_size)]

    # Add overlap between adjacent chunks
    if chunk_overlap <= 0:
        return [p for p in parts if p.strip()]

    overlapped = []
    for i, part in enumerate(parts):
        if i == 0:
            overlapped.append(part)
        else:
            # Prepend tail of previous chunk
            prev = parts[i - 1]
            tail = prev[-chunk_overlap:] if len(prev) > chunk_overlap else prev
            combined = tail + " " + part
            overlapped.append(combined.strip())

    return [c for c in overlapped if c.strip()]


# ─── Page-Aware Chunking ──────────────────────────────────────────────────────

def chunk_document(
    document_id: str,
    document_name: str,
    full_text: str,
    pages: List[str],
    chunk_size: int = 800,
    chunk_overlap: int = 120,
    category: str = "general",
    tags: Optional[List[str]] = None,
) -> List[DocumentChunk]:
    """
    Chunk a document into overlapping segments with metadata.

    Returns a list of DocumentChunk objects ready for embedding.
    """
    full_text = clean_text(full_text)
    chunks_raw = recursive_chunk(full_text, chunk_size, chunk_overlap)

    # Build a simple page mapping: character offset → page number
    page_offsets = []
    offset = 0
    for i, page in enumerate(pages):
        page_offsets.append((offset, offset + len(page), i + 1))
        offset += len(page) + 2  # +2 for "\n\n"

    def find_page(chunk_text: str) -> Optional[int]:
        pos = full_text.find(chunk_text[:80])  # match first 80 chars
        if pos == -1:
            return None
        for start, end, page_num in page_offsets:
            if start <= pos < end:
                return page_num
        return None

    document_chunks = []
    for idx, chunk_text in enumerate(chunks_raw):
        if not chunk_text.strip():
            continue

        page_num = find_page(chunk_text)
        chunk_id = f"{document_id}_chunk_{idx:04d}"

        document_chunks.append(
            DocumentChunk(
                chunk_id=chunk_id,
                document_id=document_id,
                document_name=document_name,
                text=chunk_text,
                chunk_index=idx,
                page_number=page_num,
                metadata={
                    "document_id": document_id,
                    "document_name": document_name,
                    "chunk_index": idx,
                    "page_number": page_num,
                    "category": category,
                    "tags": tags or [],
                    "text_length": len(chunk_text),
                },
            )
        )

    logger.info(
        f"Document '{document_name}' chunked into {len(document_chunks)} chunks "
        f"(chunk_size={chunk_size}, overlap={chunk_overlap})"
    )
    return document_chunks
